"""Document ingestion pipeline — parsing, chunking, embedding, indexing."""

from __future__ import annotations

import hashlib
import io
import re
import time
from collections.abc import Callable
from html.parser import HTMLParser as StdLibHTMLParser
from typing import Any

from eaip.knowledge.base import Chunker, DocumentParser, EmbeddingProvider, VectorStore
from eaip.knowledge.chunker import create_chunker
from eaip.knowledge.events import DocumentIngested
from eaip.knowledge.exceptions import (
    DocumentParseError,
    UnsupportedFormatError,
)
from eaip.knowledge.models import (
    DocumentChunk,
    DocumentFormat,
    IndexingStatus,
    IngestionConfig,
    IngestionResult,
    KnowledgeDocument,
)
from eaip.logging.context import get_logger


def _hash_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()[:16]


class TextParser:
    """Parses plain text content."""

    async def parse(self, content: bytes, **_kwargs: str) -> str:
        """Parse text content.

        Args:
            content: Raw bytes.

        Returns:
            Decoded text.
        """
        return content.decode("utf-8", errors="replace")


class MarkdownParser:
    """Parses Markdown content."""

    async def parse(self, content: bytes, **_kwargs: str) -> str:
        """Parse Markdown content into plain text.

        Args:
            content: Raw markdown bytes.

        Returns:
            Plain text.
        """
        text = content.decode("utf-8", errors="replace")

        text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
        text = re.sub(r"\[([^\]]*)\]\(.*?\)", r"\1", text)
        text = re.sub(r"#{1,6}\s+", "", text)
        text = re.sub(r"[*_`~]", "", text)
        text = re.sub(r"^[\s]*[-*+][\s]+", "", text, flags=re.MULTILINE)
        text = re.sub(r"^[\s]*\d+\.[\s]+", "", text, flags=re.MULTILINE)
        text = re.sub(r"\|", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


class _TextExtractor(StdLibHTMLParser):
    """HTML text extractor that skips script/style/nav/footer/header tags."""

    def __init__(self) -> None:
        super().__init__()
        self._result: list[str] = []
        self._skip = False

    def handle_starttag(self, tag: str, _attrs: list[tuple[str, str | None]]) -> None:
        if tag in ("script", "style", "nav", "footer", "header"):
            self._skip = True

    def handle_endtag(self, tag: str) -> None:
        if tag in ("script", "style", "nav", "footer", "header"):
            self._skip = False

    def handle_data(self, data: str) -> None:
        if not self._skip:
            stripped = data.strip()
            if stripped:
                self._result.append(stripped)

    def result(self) -> str:
        """Return extracted text."""
        return "\n".join(self._result)


class HTMLParser:
    """Parses HTML content."""

    async def parse(self, content: bytes, **_kwargs: str) -> str:
        """Parse HTML content into plain text.

        Args:
            content: Raw HTML bytes.

        Returns:
            Plain text.
        """
        text = content.decode("utf-8", errors="replace")
        extractor = _TextExtractor()
        extractor.feed(text)
        return extractor.result()


class PDFParser:
    """Parses PDF content using pypdf."""

    async def parse(self, content: bytes, **_kwargs: str) -> str:
        """Parse PDF content into plain text.

        Args:
            content: Raw PDF bytes.

        Returns:
            Extracted text.
        """
        try:
            from pypdf import PdfReader  # type: ignore[import-not-found]  # noqa: PLC0415

            reader = PdfReader(io.BytesIO(content))
            pages: list[str] = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            return "\n\n".join(pages)
        except ImportError:
            raise DocumentParseError("pypdf is required") from None
        except Exception as exc:
            raise DocumentParseError(f"Failed to parse PDF: {exc}") from exc


class DOCXParser:
    """Parses DOCX content using python-docx."""

    async def parse(self, content: bytes, **_kwargs: str) -> str:
        """Parse DOCX content into plain text.

        Args:
            content: Raw DOCX bytes.

        Returns:
            Extracted text.
        """
        try:
            from docx import Document  # type: ignore[import-not-found]  # noqa: PLC0415

            doc = Document(io.BytesIO(content))
            paragraphs: list[str] = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError:
            raise DocumentParseError("python-docx is required") from None
        except Exception as exc:
            raise DocumentParseError(f"Failed to parse DOCX: {exc}") from exc


_PARSER_REGISTRY: dict[DocumentFormat, DocumentParser] = {
    DocumentFormat.TXT: TextParser(),
    DocumentFormat.MARKDOWN: MarkdownParser(),
    DocumentFormat.HTML: HTMLParser(),
    DocumentFormat.PDF: PDFParser(),
    DocumentFormat.DOCX: DOCXParser(),
}


def get_parser(doc_format: DocumentFormat | str) -> DocumentParser:
    """Return the parser for a given document format.

    Args:
        doc_format: The document format.

    Returns:
        A DocumentParser implementation.

    Raises:
        UnsupportedFormatError: If the format is not supported.
    """
    if isinstance(doc_format, str):
        fmt = doc_format
        try:
            doc_format = DocumentFormat(doc_format)
        except ValueError:
            raise UnsupportedFormatError(f"Unsupported document format: {fmt}") from None
    parser = _PARSER_REGISTRY.get(doc_format)
    if parser is None:
        raise UnsupportedFormatError(f"Unsupported document format: {doc_format}")
    return parser


def register_parser(doc_format: DocumentFormat, parser: DocumentParser) -> None:
    """Register a custom parser for a document format.

    Args:
        doc_format: The document format.
        parser: The parser implementation.
    """
    _PARSER_REGISTRY[doc_format] = parser


class IngestionPipeline:
    """Orchestrates the document ingestion pipeline.

    Pipeline flow:
        1. Parse the raw document bytes into text.
        2. Split the text into chunks using the configured strategy.
        3. Generate embeddings for each chunk.
        4. Upsert the chunks into the vector store.
    """

    def __init__(
        self,
        config: IngestionConfig,
        vector_store: VectorStore,
        embedding_provider: EmbeddingProvider,
        event_publisher: Callable[[object], Any] | None = None,
    ) -> None:
        """Initialize the ingestion pipeline.

        Args:
            config: The ingestion configuration.
            vector_store: The vector store to index into.
            embedding_provider: The embedding provider.
            event_publisher: Optional callable to publish domain events.
        """
        self._config = config
        self._vector_store = vector_store
        self._embedding_provider = embedding_provider
        self._chunker: Chunker = create_chunker(config.chunking)
        self._event_publisher = event_publisher
        self._log = get_logger("eaip.knowledge.ingestion")

    async def ingest(
        self,
        document_id: str,
        content: bytes,
        doc_format: DocumentFormat | str,
        *,
        title: str = "",
        source: str = "",
        metadata: dict[str, Any] | None = None,
    ) -> IngestionResult:
        """Ingest a document into the knowledge store.

        Args:
            document_id: Unique identifier for the document.
            content: Raw document bytes.
            doc_format: The document format.
            title: Optional document title.
            source: Optional source path or URL.
            metadata: Optional metadata dictionary.

        Returns:
            An IngestionResult with status and details.
        """
        t0 = time.monotonic()
        fmt_name = doc_format.value if isinstance(doc_format, DocumentFormat) else doc_format
        self._log.info(
            "ingestion.start",
            document_id=document_id,
            format=fmt_name,
        )

        try:
            parser = get_parser(doc_format)
            text = await parser.parse(content)

            if not text.strip():
                raise DocumentParseError("Document produced no text content")

            content_hash = _hash_bytes(content) if self._config.generate_hash else ""

            chunk_metadata: dict[str, Any] = {
                "title": title,
                "source": source,
                **(metadata or {}),
            }

            chunks = await self._chunker.chunk(
                text,
                document_id=document_id,
                collection=self._config.collection,
                **{k: str(v) for k, v in chunk_metadata.items()},
            )

            if self._config.embedding.provider or self._config.embedding.model:
                texts = [c.content for c in chunks]
                embeddings = await self._embedding_provider.embed(texts)
                chunks = [
                    DocumentChunk(
                        chunk_id=c.chunk_id,
                        document_id=c.document_id,
                        collection=c.collection,
                        content=c.content,
                        content_hash=c.content_hash,
                        chunk_index=c.chunk_index,
                        embedding=embeddings[i] if i < len(embeddings) else c.embedding,
                        metadata=c.metadata,
                        created_at=c.created_at,
                    )
                    for i, c in enumerate(chunks)
                ]

            await self._vector_store.upsert_points(self._config.collection, chunks)

            fmt = (
                doc_format
                if isinstance(doc_format, DocumentFormat)
                else DocumentFormat(doc_format)
            )
            document = KnowledgeDocument(
                document_id=document_id,
                collection=self._config.collection,
                format=fmt,
                title=title,
                source=source,
                metadata=chunk_metadata,
                indexing_status=IndexingStatus.INDEXED,
                content_hash=content_hash,
                chunk_count=len(chunks),
            )

            duration_ms = (time.monotonic() - t0) * 1000

            if self._event_publisher is not None:
                event = DocumentIngested(
                    document_id=document_id,
                    collection=self._config.collection,
                    chunk_count=len(chunks),
                    duration_ms=duration_ms,
                )
                self._event_publisher(event)

            self._log.info(
                "ingestion.complete",
                document_id=document_id,
                chunks=len(chunks),
                duration_ms=round(duration_ms, 1),
            )

            return IngestionResult(
                document=document,
                chunk_count=len(chunks),
                status=IndexingStatus.INDEXED,
                duration_ms=duration_ms,
            )

        except (DocumentParseError, UnsupportedFormatError) as exc:
            duration_ms = (time.monotonic() - t0) * 1000
            self._log.error("ingestion.failed", document_id=document_id, error=str(exc))
            document = self._make_failed_document(document_id, doc_format, title, source, metadata)
            return IngestionResult(
                document=document,
                chunk_count=0,
                status=IndexingStatus.FAILED,
                duration_ms=duration_ms,
                errors=(str(exc),),
            )
        except Exception as exc:
            duration_ms = (time.monotonic() - t0) * 1000
            self._log.error("ingestion.error", document_id=document_id, error=str(exc))
            document = self._make_failed_document(document_id, doc_format, title, source, metadata)
            return IngestionResult(
                document=document,
                chunk_count=0,
                status=IndexingStatus.FAILED,
                duration_ms=duration_ms,
                errors=(str(exc),),
            )

    def _make_failed_document(
        self,
        document_id: str,
        doc_format: DocumentFormat | str,
        title: str,
        source: str,
        metadata: dict[str, Any] | None,
    ) -> KnowledgeDocument:
        try:
            doc_format_enum: DocumentFormat | str = doc_format
            if isinstance(doc_format, str):
                doc_format_enum = DocumentFormat(doc_format)
            return KnowledgeDocument(
                document_id=document_id,
                collection=self._config.collection,
                format=doc_format_enum,
                title=title,
                source=source,
                metadata=metadata or {},
                indexing_status=IndexingStatus.FAILED,
            )
        except Exception:
            return KnowledgeDocument(
                document_id=document_id,
                collection=self._config.collection,
                format=DocumentFormat.TXT,
                title=title,
                source=source,
                metadata=metadata or {},
                indexing_status=IndexingStatus.FAILED,
            )


__all__ = [
    "DOCXParser",
    "HTMLParser",
    "IngestionPipeline",
    "MarkdownParser",
    "PDFParser",
    "TextParser",
    "get_parser",
    "register_parser",
]
